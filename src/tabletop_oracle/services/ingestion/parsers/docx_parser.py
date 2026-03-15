"""DOCX document parser using python-docx.

Extracts text preserving heading styles, tables, and image references
from Microsoft Word documents.
"""

from __future__ import annotations

import logging
from typing import TYPE_CHECKING

from docx import Document as DocxDocumentFactory
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.table import Table as DocxTable

if TYPE_CHECKING:
    from pathlib import Path

    from docx.document import Document as DocxDocument

from tabletop_oracle.services.ingestion.models import (
    ExtractedImage,
    ParseResult,
    Section,
    Table,
)
from tabletop_oracle.services.ingestion.parsers.base import DocumentParser

logger = logging.getLogger(__name__)


class DocxParser(DocumentParser):
    """Extracts structured content from DOCX (Microsoft Word) documents.

    Uses python-docx to read heading styles for section hierarchy,
    extract tables as structured data, and capture image references.
    """

    _SUPPORTED_TYPES = frozenset(
        {
            "docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
    )

    def supports(self, content_type: str) -> bool:
        """Check if this parser handles the given content type."""
        return content_type.lower() in self._SUPPORTED_TYPES

    def parse(self, file_path: Path, metadata: dict[str, object]) -> ParseResult:
        """Parse a DOCX file into structured content.

        Args:
            file_path: Path to the DOCX file.
            metadata: Additional context (unused by DOCX parser).

        Returns:
            ParseResult with heading-based sections, tables, and image refs.
        """
        doc = DocxDocumentFactory(str(file_path))

        raw_text_parts: list[str] = []
        sections = self._build_sections(doc, raw_text_parts)
        tables = self._extract_tables(doc)
        images = self._extract_images(doc)
        raw_text = "\n".join(raw_text_parts)

        return ParseResult(
            raw_text=raw_text,
            sections=sections,
            tables=tables,
            images=images,
            metadata={
                "format": "docx",
                "word_count": len(raw_text.split()),
                "paragraph_count": len(doc.paragraphs),
            },
        )

    def _build_sections(self, doc: DocxDocument, raw_text_parts: list[str]) -> list[Section]:
        """Build sections from DOCX paragraph heading styles.

        Args:
            doc: python-docx Document object.
            raw_text_parts: Mutable list to accumulate raw text (side effect).

        Returns:
            Flat list of Section objects with levels from heading styles.
        """
        sections: list[Section] = []
        current_title = ""
        current_level = 1
        content_parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            raw_text_parts.append(para.text)

            heading_level = self._get_heading_level(para)

            if heading_level is not None:
                # Flush previous section
                if current_title or content_parts:
                    sections.append(
                        Section(
                            title=current_title,
                            level=current_level,
                            content="\n".join(content_parts).strip(),
                            source_location=f"docx:heading{current_level}",
                        )
                    )
                    content_parts = []

                current_title = text
                current_level = heading_level
            elif text:
                content_parts.append(text)

        # Flush final section
        if current_title or content_parts:
            sections.append(
                Section(
                    title=current_title,
                    level=current_level,
                    content="\n".join(content_parts).strip(),
                    source_location=f"docx:heading{current_level}",
                )
            )

        # If no sections found, create a single root section
        if not sections:
            all_text = "\n".join(p.text for p in doc.paragraphs).strip()
            sections.append(
                Section(title="", level=1, content=all_text, source_location="docx:body")
            )

        return sections

    def _get_heading_level(self, paragraph: object) -> int | None:
        """Extract heading level from a paragraph's style name.

        Args:
            paragraph: A python-docx Paragraph object.

        Returns:
            Heading level (1-9) or None if not a heading.
        """
        style = getattr(paragraph, "style", None)
        if style is None:
            return None

        style_name = getattr(style, "name", "") or ""

        # python-docx uses style names like "Heading 1", "Heading 2", etc.
        if style_name.startswith("Heading "):
            try:
                return int(style_name.split()[-1])
            except (ValueError, IndexError):
                return None

        return None

    def _extract_tables(self, doc: DocxDocument) -> list[Table]:
        """Extract tables from the DOCX document.

        Args:
            doc: python-docx Document object.

        Returns:
            List of Table objects.
        """
        tables: list[Table] = []

        for table_idx, docx_table in enumerate(doc.tables):
            if not isinstance(docx_table, DocxTable):
                continue

            rows_data = [[cell.text.strip() for cell in row.cells] for row in docx_table.rows]

            if len(rows_data) < 1:
                continue

            # First row as headers
            headers = rows_data[0]
            rows = [
                {headers[i]: (row[i] if i < len(row) else "") for i in range(len(headers))}
                for row in rows_data[1:]
            ]

            tables.append(
                Table(
                    headers=headers,
                    rows=rows,
                    caption=f"Table {table_idx + 1}",
                )
            )

        return tables

    def _extract_images(self, doc: DocxDocument) -> list[ExtractedImage]:
        """Extract image references from the DOCX document.

        Uses the document's relationship graph to find embedded images.

        Args:
            doc: python-docx Document object.

        Returns:
            List of ExtractedImage references.
        """
        images: list[ExtractedImage] = []
        img_idx = 0

        for rel in doc.part.rels.values():
            if rel.reltype == RT.IMAGE:
                img_idx += 1
                target = rel.target_ref

                # Determine format from the target path
                fmt = "unknown"
                lower_target = str(target).lower()
                for ext in ("png", "jpg", "jpeg", "gif", "tiff", "bmp"):
                    if lower_target.endswith(f".{ext}"):
                        fmt = "jpeg" if ext == "jpg" else ext
                        break

                images.append(
                    ExtractedImage(
                        image_data=b"",  # Reference only
                        format=fmt,
                        location=f"docx:image:{target}",
                        alt_text=f"Image {img_idx}",
                    )
                )

        return images
